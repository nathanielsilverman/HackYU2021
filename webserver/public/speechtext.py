import speech_recognition as sr
from deep_translator import GoogleTranslator
import time
import datetime

import docx
# doc = docx.Document("/worddoc.docx")

# from emoji_translate.emoji_translate import Translator

# emo = Translator(exact_match_only=False, randomize=True)
dest_lang = input('What language do you want your lecture in?')
file_name = input("What file do you want this to be written to?")+".docx"
to_translate = 'I want to translate this text'

start_time = datetime.datetime(100, 1, 1, 0, 0, 0)
max_time = datetime.datetime(100, 1, 1, 0, 0, 15)
start = time.time()

mydoc = docx.Document()
recognizer = sr.Recognizer()
# for i in range(2):
x = 3
while True:
    while time.time() - start <= x:
        with sr.Microphone() as source:
            print("Talk")
            audio_text = recognizer.record(source, duration=3)
            # audio_text = recognizer.listen(source)
            st = recognizer.recognize_google(audio_text)
            # print(type(st))
            print("Time over, thanks")

            try:
                translated = GoogleTranslator(source='auto', target=dest_lang).translate(st)
                mydoc.add_paragraph(translated)
                mydoc.save(file_name)
                # print(emo.emojify(text))
            except:
                print("Sorry, I did not get that")
        x = x + 3